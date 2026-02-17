"""
DOCX Generation Service – 3 visually distinct templates.
Professional: Calibri, centered, standard corporate look.
Modern: Left-aligned, navy/green accents, FAANG-style.
Classic: Times New Roman, serif, ruled dividers.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from backend.app.models.schemas import ResumeContent
from io import BytesIO


# ── Brand colours ──────────────────────────────────────────
NAVY  = RGBColor(0x1E, 0x3A, 0x4C)
GREEN = RGBColor(0x2E, 0x7D, 0x52)
BLACK = RGBColor(0x11, 0x11, 0x11)
DGRAY = RGBColor(0x44, 0x44, 0x44)
LGRAY = RGBColor(0x88, 0x88, 0x88)


def _add_bottom_border(paragraph, color_hex="000000"):
    """Add bottom border (underline) to a paragraph via XML."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


class DOCXGenerationService:

    def generate_docx(self, resume_content: ResumeContent, template: str = "professional") -> BytesIO:
        doc = Document()
        for sec in doc.sections:
            sec.top_margin    = Inches(0.7)
            sec.bottom_margin = Inches(0.7)
            sec.left_margin   = Inches(0.75)
            sec.right_margin  = Inches(0.75)

        if template == "modern":
            self._modern(doc, resume_content)
        elif template == "classic":
            self._classic(doc, resume_content)
        else:
            self._professional(doc, resume_content)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    # ═══════════════════════════════════════════════════════
    #  PROFESSIONAL – Calibri, centered header, thin dividers
    # ═══════════════════════════════════════════════════════
    def _professional(self, doc: Document, rc: ResumeContent):

        def name_para():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(rc.full_name)
            r.font.name = 'Calibri'; r.font.size = Pt(22); r.bold = True
            r.font.color.rgb = BLACK
            return p

        def contact_para():
            parts = _cp(rc)
            if not parts: return
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(' | '.join(parts))
            r.font.name = 'Calibri'; r.font.size = Pt(9.5)
            r.font.color.rgb = DGRAY

        def sec_hdr(title):
            p = doc.add_paragraph()
            _add_bottom_border(p, "AAAAAA")
            r = p.add_run(title.upper())
            r.font.name = 'Calibri'; r.font.size = Pt(11); r.bold = True
            r.font.color.rgb = BLACK
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(4)

        def job_title(text):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = 'Calibri'; r.font.size = Pt(10.5); r.bold = True
            r.font.color.rgb = BLACK
            p.paragraph_format.space_after = Pt(1)

        def sub_line(text):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = 'Calibri'; r.font.size = Pt(9.5); r.italic = True
            r.font.color.rgb = LGRAY
            p.paragraph_format.space_after = Pt(3)

        def bullet(text):
            p = doc.add_paragraph(style='List Bullet')
            r = p.add_run(text)
            r.font.name = 'Calibri'; r.font.size = Pt(9.5)
            r.font.color.rgb = BLACK
            p.paragraph_format.left_indent  = Inches(0.2)
            p.paragraph_format.space_after  = Pt(2)

        def body(text):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = 'Calibri'; r.font.size = Pt(10)
            r.font.color.rgb = BLACK
            p.paragraph_format.space_after = Pt(5)

        # ── Build ──
        name_para()
        contact_para()
        doc.add_paragraph()

        if rc.summary: sec_hdr('Professional Summary'); body(rc.summary)
        if rc.skills:  sec_hdr('Skills'); body(', '.join(rc.skills))

        if rc.experience:
            sec_hdr('Professional Experience')
            for exp in rc.experience:
                job_title(exp.title)
                s = f'{exp.company} | {exp.duration}'
                if exp.location: s += f' | {exp.location}'
                sub_line(s)
                for r in (exp.responsibilities or []): bullet(r)
                doc.add_paragraph()

        if rc.education:
            sec_hdr('Education')
            for edu in rc.education:
                job_title(edu.degree)
                s = f'{edu.institution} | {edu.year}'
                if edu.gpa: s += f' | GPA: {edu.gpa}'
                sub_line(s)

        if rc.projects:
            sec_hdr('Projects')
            for p in rc.projects:
                job_title(p.name)
                if p.technologies: sub_line(f'Technologies: {p.technologies}')
                if p.description:  bullet(p.description)
                if p.impact:       bullet(p.impact)
                doc.add_paragraph()

        if rc.certifications:
            sec_hdr('Certifications')
            for c in rc.certifications:
                if c and c.strip(): bullet(c)

        if rc.achievements:
            sec_hdr('Achievements')
            for a in rc.achievements:
                if a and a.strip(): bullet(a)

    # ═══════════════════════════════════════════════════════
    #  MODERN – Calibri, navy name, green section headers
    # ═══════════════════════════════════════════════════════
    def _modern(self, doc: Document, rc: ResumeContent):

        def name_para():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(rc.full_name)
            r.font.name = 'Calibri'; r.font.size = Pt(26); r.bold = True
            r.font.color.rgb = NAVY
            return p

        def accent_line():
            p = doc.add_paragraph()
            _add_bottom_border(p, "2E7D52")
            p.paragraph_format.space_after = Pt(4)

        def contact_para():
            parts = _cp(rc)
            if not parts: return
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run('  ·  '.join(parts))
            r.font.name = 'Calibri'; r.font.size = Pt(9.5)
            r.font.color.rgb = DGRAY

        def sec_hdr(title):
            p = doc.add_paragraph()
            r = p.add_run(title.upper())
            r.font.name = 'Calibri'; r.font.size = Pt(11); r.bold = True
            r.font.color.rgb = GREEN
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(3)
            _add_bottom_border(p, "D1E8DB")

        def job_title(text):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = 'Calibri'; r.font.size = Pt(11); r.bold = True
            r.font.color.rgb = NAVY
            p.paragraph_format.space_after = Pt(1)

        def sub_line(text):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = 'Calibri'; r.font.size = Pt(9.5)
            r.font.color.rgb = LGRAY
            p.paragraph_format.space_after = Pt(3)

        def bullet(text, symbol='▸'):
            p = doc.add_paragraph()
            r = p.add_run(f'{symbol}  {text}')
            r.font.name = 'Calibri'; r.font.size = Pt(9.5)
            r.font.color.rgb = BLACK
            p.paragraph_format.left_indent  = Inches(0.18)
            p.paragraph_format.space_after  = Pt(2)

        def body(text):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = 'Calibri'; r.font.size = Pt(10)
            r.font.color.rgb = BLACK
            p.paragraph_format.space_after = Pt(4)

        # ── Build ──
        name_para()
        accent_line()
        contact_para()
        doc.add_paragraph()

        if rc.summary: sec_hdr('Profile'); body(rc.summary)
        if rc.skills:  sec_hdr('Core Skills'); body('   ▸   '.join(rc.skills))

        if rc.experience:
            sec_hdr('Experience')
            for exp in rc.experience:
                job_title(exp.title)
                s = f'{exp.company}'
                if exp.location: s += f', {exp.location}'
                s += f'  |  {exp.duration}'
                sub_line(s)
                for r in (exp.responsibilities or []): bullet(r)
                doc.add_paragraph()

        if rc.education:
            sec_hdr('Education')
            for edu in rc.education:
                job_title(edu.degree)
                s = f'{edu.institution}  |  {edu.year}'
                if edu.gpa: s += f'  |  CGPA: {edu.gpa}'
                sub_line(s)

        if rc.projects:
            sec_hdr('Projects')
            for p in rc.projects:
                job_title(p.name)
                if p.technologies: sub_line(f'Stack: {p.technologies}')
                if p.description:  bullet(p.description)
                if p.impact:       bullet(f'Impact: {p.impact}')
                doc.add_paragraph()

        if rc.certifications:
            sec_hdr('Certifications')
            for c in rc.certifications:
                if c and c.strip(): bullet(c)

        if rc.achievements:
            sec_hdr('Achievements & Awards')
            for a in rc.achievements:
                if a and a.strip(): bullet(a)

    # ═══════════════════════════════════════════════════════
    #  CLASSIC – Times New Roman, serif, ruled dividers
    # ═══════════════════════════════════════════════════════
    def _classic(self, doc: Document, rc: ResumeContent):

        def name_para():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(rc.full_name)
            r.font.name = 'Times New Roman'; r.font.size = Pt(24); r.bold = True
            r.font.color.rgb = BLACK

        def full_rule():
            p = doc.add_paragraph()
            _add_bottom_border(p, "111111")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)

        def contact_para():
            parts = _cp(rc)
            if not parts: return
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(' | '.join(parts))
            r.font.name = 'Times New Roman'; r.font.size = Pt(10)
            r.font.color.rgb = DGRAY

        def sec_hdr(title):
            full_rule()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(title)
            r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
            r.font.color.rgb = BLACK
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            full_rule()

        def job_title(text):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
            r.font.color.rgb = BLACK
            p.paragraph_format.space_after = Pt(1)

        def sub_line(text):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
            r.font.color.rgb = DGRAY
            p.paragraph_format.space_after = Pt(3)

        def bullet(text):
            p = doc.add_paragraph()
            r = p.add_run(f'•   {text}')
            r.font.name = 'Times New Roman'; r.font.size = Pt(10)
            r.font.color.rgb = BLACK
            p.paragraph_format.left_indent  = Inches(0.2)
            p.paragraph_format.space_after  = Pt(3)

        def body(text):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = 'Times New Roman'; r.font.size = Pt(10.5)
            r.font.color.rgb = BLACK
            p.paragraph_format.space_after = Pt(5)

        # ── Build ──
        name_para()
        full_rule()
        contact_para()
        doc.add_paragraph()

        if rc.summary: sec_hdr('PROFESSIONAL SUMMARY'); body(rc.summary)
        if rc.skills:  sec_hdr('SKILLS'); body(' • '.join(rc.skills))

        if rc.experience:
            sec_hdr('PROFESSIONAL EXPERIENCE')
            for exp in rc.experience:
                job_title(exp.title)
                s = f'{exp.company}  |  {exp.duration}'
                if exp.location: s += f'  |  {exp.location}'
                sub_line(s)
                for r in (exp.responsibilities or []): bullet(r)
                doc.add_paragraph()

        if rc.education:
            sec_hdr('EDUCATION')
            for edu in rc.education:
                job_title(edu.degree)
                s = f'{edu.institution}  |  {edu.year}'
                if edu.gpa: s += f'  |  GPA: {edu.gpa}'
                sub_line(s)

        if rc.projects:
            sec_hdr('PROJECTS')
            for p in rc.projects:
                job_title(p.name)
                if p.technologies: sub_line(f'Technologies: {p.technologies}')
                if p.description:  bullet(p.description)
                if p.impact:       bullet(p.impact)
                doc.add_paragraph()

        if rc.certifications:
            sec_hdr('CERTIFICATIONS')
            for c in rc.certifications:
                if c and c.strip(): bullet(c)

        if rc.achievements:
            sec_hdr('ACHIEVEMENTS')
            for a in rc.achievements:
                if a and a.strip(): bullet(a)


# ── Shared helper ──────────────────────────────────────────
def _cp(rc: ResumeContent):
    parts = []
    if rc.contact.phone:    parts.append(rc.contact.phone)
    if rc.contact.email:    parts.append(rc.contact.email)
    if rc.contact.linkedin: parts.append(rc.contact.linkedin)
    if rc.contact.github:   parts.append(rc.contact.github)
    return parts