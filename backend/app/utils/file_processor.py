"""
File Processing Utilities
Handles PDF and DOCX file upload processing
"""
from pypdf import PdfReader
from docx import Document
import io
from typing import Optional
import json
import re
from google.genai import types
from backend.app.services.gemini_service import GeminiService


class FileProcessor:
    """Process uploaded resume files"""
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_content: PDF file content as bytes
            
        Returns:
            Extracted text as string
        """
        try:
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)
            
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
            
        except Exception as e:
            print(f"Error extracting PDF text: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_content: DOCX file content as bytes
            
        Returns:
            Extracted text as string
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            
            return text.strip()
            
        except Exception as e:
            print(f"Error extracting DOCX text: {str(e)}")
            return ""
    
    @staticmethod
    def process_uploaded_file(filename: str, file_content: bytes) -> Optional[str]:
        """
        Process uploaded file and extract text
        
        Args:
            filename: Name of uploaded file
            file_content: File content as bytes
            
        Returns:
            Extracted text or None if processing fails
        """
        file_extension = filename.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            return FileProcessor.extract_text_from_pdf(file_content)
        elif file_extension in ['docx', 'doc']:
            return FileProcessor.extract_text_from_docx(file_content)
        else:
            return None
    
    @staticmethod
    def validate_file(filename: str, file_size: int, max_size: int = 10 * 1024 * 1024) -> tuple[bool, str]:
        """
        Validate uploaded file
        
        Args:
            filename: Name of uploaded file
            file_size: Size of file in bytes
            max_size: Maximum allowed file size in bytes
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check file extension
        allowed_extensions = {'.pdf', '.docx', '.doc'}
        file_extension = '.' + filename.lower().split('.')[-1]
        
        if file_extension not in allowed_extensions:
            return False, f"Invalid file type. Allowed types: {', '.join(allowed_extensions)}"
        
        # Check file size
        if file_size > max_size:
            max_size_mb = max_size / (1024 * 1024)
            return False, f"File size exceeds maximum allowed size of {max_size_mb}MB"
        
        return True, ""

    @staticmethod
    def parse_resume_to_structured_data(text: str, full_name: str, target_role: str) -> dict:
        """
        Parse raw resume text into structured user_data format using Gemini
        Returns same format as manual form for consistent processing
        """
        try:
            gemini = GeminiService()
            prompt = f"""
            PARSE this resume text into structured JSON format EXACTLY matching this schema:
            
            Name: {full_name}
            Target Role: {target_role}
            
            Resume text:
            {text[:4000]}
            
            Return ONLY JSON with this EXACT structure:
            {{
                "full_name": "{full_name}",
                "summary": "Extracted/created summary",
                "tech_skills": ["Python", "AWS", "Docker"],
                "soft_skills": ["Leadership", "Communication"],
                "experience": [{{
                    "title": "Software Engineer",
                    "company": "Google", 
                    "duration": "2022-Present",
                    "location": "Bangalore",
                    "responsibilities": ["Built scalable APIs", "Led team of 5"]
                }}],
                "education": [{{
                    "degree": "B.Tech Computer Science",
                    "institution": "IIT Madras",
                    "year": "2022",
                    "gpa": "8.5"
                }}],
                "projects": [{{
                    "name": "Smart Expense Tracker",
                    "description": "Full-stack expense management app",
                    "technologies": "React, Node.js, MongoDB",
                    "impact": "Deployed to 500+ users"
                }}],
                "certifications": ["AWS Certified Solutions Architect", "Google Data Analytics"],
                "achievements": ["Won hackathon 2023", "Top 5% LeetCode"]
            }}
            
            Extract ALL sections. If missing, use reasonable defaults for {target_role}.
            """
            
            # Use same Gemini service to parse
            config = types.GenerateContentConfig(
                temperature=0.1,  # Very low for structured output
                response_mime_type="application/json"
            )
            
            response = gemini.client.models.generate_content(
                model=gemini.model_id,
                contents=prompt,
                config=config
            )
            
            raw_json = response.text
            cleaned = re.sub(r'^```json\s*|\s*```$', '', raw_json.strip(), flags=re.MULTILINE)
            
            structured_data = json.loads(cleaned)
            print(f"📦 Parsed resume → {len(json.dumps(structured_data))} chars structured data")
            
            return structured_data
            
        except Exception as e:
            print(f"⚠️ Resume parsing failed: {e}")
            # Fallback: minimal structured data
            return {
                "full_name": full_name,
                "summary": "Experienced professional seeking opportunities",
                "tech_skills": [],
                "soft_skills": [],
                "experience": [],
                "education": [],
                "projects": [],
                "certifications": [],
                "achievements": []
            }
